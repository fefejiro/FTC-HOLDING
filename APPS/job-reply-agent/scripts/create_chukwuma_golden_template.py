from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "instances" / "chukwuma" / "resumes" / "Chukwuma Mezie-Okoye Golden Template.docx"
ORANGE = "F4511E"
CHARCOAL = "20252B"
MUTED = "68717A"
LIGHT = "D9DEE3"


def set_cell_width(cell, width_twips):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def set_bottom_border(cell, color=LIGHT, size=6):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    bottom = borders.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        borders.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:color"), color)


def style_run(run, size=9.5, bold=False, color=CHARCOAL):
    run.font.name = "Lato"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Lato")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def paragraph(cell, text="", size=9.5, bold=False, color=CHARCOAL, before=0, after=3):
    p = cell.add_paragraph() if cell.paragraphs[0].text else cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    style_run(p.add_run(text), size=size, bold=bold, color=color)
    return p


def bullet(cell, text):
    p = cell.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.22)
    p.paragraph_format.first_line_indent = Inches(-0.12)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.0
    style_run(p.add_run(text), size=9.15)
    return p


def add_section(table, label, writer):
    row = table.add_row()
    left, right = row.cells
    left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    paragraph(left, label, size=10.5, bold=True)
    writer(right)
    set_bottom_border(left)
    set_bottom_border(right)
    set_table_geometry(table, [2700, 6660])


def add_role(cell, title, employer, dates, bullets):
    p = cell.add_paragraph() if cell.paragraphs[0].text else cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(0)
    style_run(p.add_run(f"{title} | {employer}"), size=9.6, bold=True)
    d = cell.add_paragraph()
    d.paragraph_format.space_after = Pt(2)
    style_run(d.add_run(dates), size=8.7, color=MUTED)
    for item in bullets:
        bullet(cell, item)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.42)
section.bottom_margin = Inches(0.42)
section.left_margin = Inches(0.48)
section.right_margin = Inches(0.48)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Lato"
normal.font.size = Pt(9.5)
normal.paragraph_format.space_after = Pt(3)

header = doc.add_table(rows=1, cols=2)
header.style = "Table Grid"
header.autofit = False
left, right = header.rows[0].cells
set_cell_margins(left, 90, 100, 100, 120)
set_cell_margins(right, 90, 130, 100, 100)
paragraph(left, "DIGITAL", size=12.5, bold=True, color=ORANGE)
paragraph(left, "TRANSFORMATION", size=12.5, bold=True, color=ORANGE)
paragraph(left, "LEADER", size=12.5, bold=True, color=ORANGE)
paragraph(
    left,
    "Communications | Product | Growth | AI Adoption",
    size=8.2,
    bold=True,
    color=MUTED,
    before=4,
)
paragraph(right, "Chukwuma Mezie-Okoye", size=14, bold=True)
paragraph(right, "Pickering, ON  L1V 0E9", size=9)
paragraph(right, "+234 802 784 4846", size=9, color=ORANGE)
paragraph(right, "chukwumamezok@gmail.com", size=9)
paragraph(right, "linkedin.com/in/chukwuma-mezie-okoye", size=9)
set_table_geometry(header, [2700, 6660])
set_bottom_border(left, color=CHARCOAL, size=14)
set_bottom_border(right, color=CHARCOAL, size=14)

table = doc.add_table(rows=0, cols=2)
table.style = "Table Grid"
table.autofit = False
set_table_geometry(table, [2700, 6660])


def write_summary(cell):
    paragraph(
        cell,
        "Digital transformation, communications, and growth leader with more than a decade of experience building brands, launching media and technology ventures, leading corporate communications, and delivering measurable audience, engagement, lead-generation, and product outcomes.",
        size=9.5,
    )
    bullet(cell, "Combines corporate communications, brand strategy, product development, digital project delivery, and stakeholder leadership.")
    bullet(cell, "Built and scaled ventures spanning media, technology-enabled marketing, e-commerce automation, digital skills, and startup incubation.")
    bullet(cell, "Experienced with AI adoption, no-code workflows, training, ecosystem development, and emerging-technology programmes.")


add_section(table, "Summary", write_summary)


def write_skills(cell):
    paragraph(cell, "Leadership & Transformation", size=9.5, bold=True, color=ORANGE)
    bullet(cell, "Digital transformation, product development, project delivery, business development, partnerships, and ecosystem building")
    paragraph(cell, "Communications & Growth", size=9.5, bold=True, color=ORANGE, before=3)
    bullet(cell, "Corporate communications, internal communications, PR, crisis communications, executive communications, brand strategy, content, SEO, and growth marketing")
    paragraph(cell, "Technology & Delivery", size=9.5, bold=True, color=ORANGE, before=3)
    bullet(cell, "AI adoption, no-code machine learning, n8n, KNIME, WordPress, Adobe tools, analytics, e-commerce automation, training, and facilitation")


add_section(table, "Skills", write_skills)


def write_experience(cell):
    add_role(
        cell,
        "Founder & CEO",
        "Activ8 Hub / Activ8 Hybrid Limited",
        "Current",
        [
            "Founded and leads an innovation hub delivering digital-skills training, startup incubation, mentorship, and creative development.",
            "Designs programmes covering AI adoption, Web3 literacy, digital inclusion, and the creator economy.",
            "Secured participation in the UNDP Young Africa Innovates Programme in 2025 and develops grant proposals and international partnerships.",
        ],
    )
    add_role(
        cell,
        "Group Director, Technology, Media & Marketing",
        "Bonitas Group",
        "2015 - Present",
        [
            "Built media, communications, and technology businesses covering corporate storytelling, PR, digital transformation, B2C product development, and e-commerce automation.",
            "Scaled WokeNationTV to more than 50,000 followers, achieved AdSense monetisation within one year, and grew website traffic to 10,000 monthly visits.",
            "Built a digital marketing system associated with a 40% increase in sales leads, 15% lift in conversions, and sustained 25% increase in brand recognition.",
            "Delivered more than 200 projects and supported Anambra State Government technology deployment across 22 Smart Schools.",
            "Built GoDigital, an AI-enabled digital communications product with packaged services for individuals and organisations.",
        ],
    )
    add_role(
        cell,
        "Brand & Communications Consultant",
        "Ventrae Integrated Services Limited & Affiliates",
        "Aug 2021 - Sep 2022",
        [
            "Led rebranding programmes associated with a 20% rise in brand recognition and 15% increase in customer engagement.",
            "Launched brand websites that increased traffic by 25% and lead generation by 30%.",
            "Managed more than 10 media, PR, creative, digital, and government stakeholder relationships.",
        ],
    )
    add_role(
        cell,
        "Head of Corporate Communications",
        "Nestoil Group / Obijackson Group",
        "Nov 2018 - Apr 2021",
        [
            "Led corporate communications, PR, digital strategy, crisis management, executive communications, and the 2019 group-wide rebrand.",
            "Deployed crisis communication and social-listening tools associated with a 50% reduction in negative-news spread and more than 80% positive host-community media mentions.",
            "Delivered a 50% increase in industry media coverage and 25% more executive speaking engagements through structured thought leadership.",
            "Managed corporate social media, internal communications, quarterly magazine production, and executive presentations.",
        ],
    )
    add_role(
        cell,
        "Client Service Executive / Project Manager",
        "Wild Fusion",
        "Feb 2017 - Jan 2018",
        [
            "Supported strategy, client pitching, campaign optimisation, digital solutions, lead generation, and account retention.",
        ],
    )
    add_role(
        cell,
        "Brand & Creative Officer",
        "Employee Energy",
        "Aug 2014 - Jan 2017",
        [
            "Designed and implemented First Bank Insurance's Competency Management Framework and facilitated national training programmes.",
        ],
    )
    add_role(
        cell,
        "Brand & Marketing Officer",
        "Brandvertisement Concepts",
        "Aug 2012 - Apr 2014",
        [
            "Delivered brand identity, marketing communications, out-of-home advertising management, and media partnerships.",
        ],
    )


add_section(table, "Experience", write_experience)


def write_education(cell):
    paragraph(cell, "B.A., English Studies & Literature", size=9.5, bold=True)
    paragraph(cell, "University of Port Harcourt, Nigeria | 2011", size=9, color=MUTED)


add_section(table, "Education", write_education)


def write_certifications(cell):
    for item in [
        "No-Code AI & Machine Learning - MIT Professional Education / Great Learning (in progress, 2026)",
        "UNDP Young Africa Innovates Programme (2025)",
        "Certificate of Competence in Grant Writing (2024)",
        "Accelerated Leadership & Business, The Edge Business School (2024)",
        "Product Management Course (2024)",
        "Digital Disruption Masterclass, Agile & Hackathons (2023)",
        "Associate Registered Practitioner in Advertising, APCON (2020)",
        "ISO 9001:2015 Quality Management Auditor (2020)",
        "Google Digital Skills for Africa (2017 and 2020)",
    ]:
        bullet(cell, item)


add_section(table, "Certifications", write_certifications)

doc.core_properties.title = "Chukwuma Mezie-Okoye Golden Resume Template"
doc.core_properties.subject = "Una Labs JobAgent approved source template"
doc.core_properties.author = "Chukwuma Mezie-Okoye"
OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
